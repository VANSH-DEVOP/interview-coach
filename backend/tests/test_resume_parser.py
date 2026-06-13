"""Tests for resume text extraction."""

import pytest
from io import BytesIO
from pypdf import PdfWriter
from docx import Document

from app.services.resume_parser import ResumeParser


@pytest.fixture
def sample_pdf():
    """Create a simple PDF with test content."""
    pdf_writer = PdfWriter()
    pdf_writer.add_blank_page(width=200, height=200)
    page = pdf_writer.pages[0]
    # Add text to the page
    from pypdf._page import PageObject, Transformation
    page.merge_page(PdfWriter().add_blank_page(width=200, height=200).pages[0])
    
    # For testing, we'll use a minimal PDF structure
    output = BytesIO()
    pdf_writer.write(output)
    output.seek(0)
    return output.getvalue()


@pytest.fixture
def sample_docx():
    """Create a simple DOCX with test content."""
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("Email: john@example.com")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def test_parse_docx(sample_docx):
    """Test DOCX parsing extracts text correctly."""
    parsed = ResumeParser.parse_docx(sample_docx)
    assert "John Doe" in parsed
    assert "Senior Software Engineer" in parsed
    assert "john@example.com" in parsed
    assert "Skills" in parsed


def test_parse_pdf_raises_with_invalid_content():
    """Test PDF parsing fails gracefully with invalid content."""
    with pytest.raises(ValueError, match="Failed to parse PDF"):
        ResumeParser.parse_pdf(b"not a pdf")


def test_parse_docx_raises_with_invalid_content():
    """Test DOCX parsing fails gracefully with invalid content."""
    with pytest.raises(ValueError, match="Failed to parse DOCX"):
        ResumeParser.parse_docx(b"not a docx")


def test_parse_dispatches_by_content_type(sample_docx):
    """Test parse() dispatches to correct parser based on content type."""
    result = ResumeParser.parse(
        sample_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "John Doe" in result


def test_parse_raises_with_unsupported_content_type():
    """Test parse() rejects unsupported content types."""
    with pytest.raises(ValueError, match="Unsupported content type"):
        ResumeParser.parse(b"data", "application/json")
